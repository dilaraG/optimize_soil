#!/usr/bin/env python3
# -*- coding: utf-8 -*-
"""Convert VKR markdown chapters to Word (.docx) with basic GOST-like typography."""

from __future__ import annotations

import argparse
import re
import sys
from pathlib import Path

import mistune
from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_BREAK, WD_LINE_SPACING
from docx.oxml.ns import qn
from docx.shared import Cm, Pt
from mistune.plugins.math import math
from mistune.plugins.table import table


def _latex_to_dollar_math(text: str) -> str:
    """Переводит \\(...\\) и \\[...\\] в $...$ / $$...$$ для mistune math."""
    text = re.sub(
        r"\\\[\s*\n([\s\S]*?)\n\s*\\\]",
        lambda m: "$$\n" + m.group(1).strip("\n") + "\n$$",
        text,
    )
    text = re.sub(
        r"\\\((?:[^\\]|\\.)*?\\\)",
        lambda m: "$" + m.group(0)[2:-2] + "$",
        text,
    )
    return text


def _set_run_font(run, name: str = "Times New Roman") -> None:
    run.font.name = name
    r = run._element.rPr
    if r is None:
        return
    rfonts = r.rFonts
    if rfonts is None:
        return
    rfonts.set(qn("w:eastAsia"), name)


def _configure_styles(doc: Document) -> None:
    normal = doc.styles["Normal"]
    normal.font.name = "Times New Roman"
    normal.font.size = Pt(14)
    pf = normal.paragraph_format
    pf.line_spacing_rule = WD_LINE_SPACING.MULTIPLE
    pf.line_spacing = 1.5
    pf.space_after = Pt(0)
    pf.first_line_indent = Cm(1.25)

    for lvl in range(1, 4):
        try:
            hs = doc.styles[f"Heading {lvl}"]
        except KeyError:
            continue
        hs.font.name = "Times New Roman"
        hs.font.bold = True
        hs.font.size = Pt(14 if lvl >= 2 else 16)
        hpf = hs.paragraph_format
        hpf.first_line_indent = Cm(0)
        hpf.space_before = Pt(12 if lvl == 1 else 10)
        hpf.space_after = Pt(6)
        hpf.line_spacing_rule = WD_LINE_SPACING.MULTIPLE
        hpf.line_spacing = 1.5


def _set_math_run(run, size_pt: float = 14) -> None:
    name = "Cambria Math"
    run.font.name = name
    r = run._element.rPr
    if r is not None:
        rfonts = r.rFonts
        if rfonts is not None:
            rfonts.set(qn("w:eastAsia"), name)
    run.font.size = Pt(size_pt)


def _clear_paragraph_runs(paragraph) -> None:
    for run in list(paragraph.runs):
        run._element.getparent().remove(run._element)


def _add_inline(paragraph, children: list | None) -> None:
    if not children:
        return
    for ch in children:
        t = ch.get("type")
        if t == "text":
            run = paragraph.add_run(ch.get("raw", ""))
            _set_run_font(run)
        elif t == "strong":
            start = len(paragraph.runs)
            _add_inline(paragraph, ch.get("children"))
            for run in paragraph.runs[start:]:
                run.bold = True
        elif t == "emphasis":
            start = len(paragraph.runs)
            _add_inline(paragraph, ch.get("children"))
            for run in paragraph.runs[start:]:
                run.italic = True
        elif t == "inline_math":
            run = paragraph.add_run(ch.get("raw", ""))
            _set_math_run(run)
        elif t == "linebreak":
            br = paragraph.add_run()
            _set_run_font(br)
            br.add_break(WD_BREAK.LINE)
        elif t == "codespan":
            run = paragraph.add_run(ch.get("raw", ""))
            _set_run_font(run, "Courier New")
            run.font.size = Pt(11)
        elif t == "link":
            inner = []
            for sub in ch.get("children") or []:
                if sub.get("type") == "text":
                    inner.append(sub.get("raw", ""))
            run = paragraph.add_run("".join(inner))
            _set_run_font(run)
        else:
            # fallback: flatten unknown inline
            raw = ch.get("raw")
            if raw:
                run = paragraph.add_run(raw)
                _set_run_font(run)
            else:
                for sub in ch.get("children") or []:
                    _add_inline(paragraph, [sub])


def _paragraph_body_style(paragraph) -> None:
    paragraph.paragraph_format.first_line_indent = Cm(1.25)
    paragraph.paragraph_format.line_spacing_rule = WD_LINE_SPACING.MULTIPLE
    paragraph.paragraph_format.line_spacing = 1.5


def _add_table(doc: Document, node: dict) -> None:
    """Строки таблицы — списки AST-узлов содержимого ячеек (для формул и выделений)."""
    rows: list[list[list]] = []

    def append_row(row_node: dict) -> None:
        cells_out: list[list] = []
        for cell in row_node.get("children") or []:
            if cell.get("type") == "table_cell":
                cells_out.append(list(cell.get("children") or []))
        if cells_out:
            rows.append(cells_out)

    for child in node.get("children") or []:
        if child.get("type") == "table_head":
            # mistune 3: под table_head лежат table_cell напрямую (одна строка заголовков),
            # а не table_row — иначе заголовок теряется в Word.
            head_ch = child.get("children") or []
            if head_ch and head_ch[0].get("type") == "table_cell":
                cells_out = [
                    list(c.get("children") or [])
                    for c in head_ch
                    if c.get("type") == "table_cell"
                ]
                if cells_out:
                    rows.append(cells_out)
            else:
                for row in head_ch:
                    if row.get("type") == "table_row":
                        append_row(row)
        elif child.get("type") == "table_body":
            for row in child.get("children") or []:
                if row.get("type") == "table_row":
                    append_row(row)

    if not rows:
        return
    ncols = max(len(r) for r in rows)
    table = doc.add_table(rows=len(rows), cols=ncols)
    table.style = "Table Grid"
    for i, r in enumerate(rows):
        for j in range(ncols):
            cell = table.rows[i].cells[j]
            p = cell.paragraphs[0]
            _clear_paragraph_runs(p)
            children = r[j] if j < len(r) else []
            _add_inline(p, children)
            for run in p.runs:
                if run.font.name != "Cambria Math":
                    _set_run_font(run)
                if run.font.size is None:
                    run.font.size = Pt(12)
                if i == 0:
                    run.bold = True
            p.paragraph_format.first_line_indent = Cm(0)
            p.paragraph_format.line_spacing_rule = WD_LINE_SPACING.SINGLE


def _list_item_inline_children(item: dict) -> list:
    merged: list = []
    for c in item.get("children") or []:
        if c.get("type") == "block_text":
            merged.extend(c.get("children") or [])
        else:
            merged.append(c)
    return merged


def _process_blocks(doc: Document, blocks: list) -> None:
    for node in blocks:
        t = node.get("type")
        if t in ("blank_line",):
            continue
        if t == "thematic_break":
            p = doc.add_paragraph()
            p.paragraph_format.first_line_indent = Cm(0)
            p.add_run("—" * 20)
            continue
        if t == "heading":
            level = int(node.get("attrs", {}).get("level", 2))
            level = min(max(level, 1), 3)
            text_parts: list[str] = []

            def collect_heading(chlist):
                for c in chlist or []:
                    if c.get("type") == "text":
                        text_parts.append(c.get("raw", ""))
                    else:
                        collect_heading(c.get("children"))

            collect_heading(node.get("children"))
            title = "".join(text_parts).strip()
            h = doc.add_heading(title, level=level)
            for run in h.runs:
                _set_run_font(run)
            h.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.LEFT
            continue
        if t == "block_math":
            p = doc.add_paragraph()
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            p.paragraph_format.first_line_indent = Cm(0)
            p.paragraph_format.space_before = Pt(6)
            p.paragraph_format.space_after = Pt(6)
            p.paragraph_format.line_spacing_rule = WD_LINE_SPACING.MULTIPLE
            p.paragraph_format.line_spacing = 1.5
            run = p.add_run((node.get("raw") or "").strip())
            _set_math_run(run, 14)
            continue
        if t == "paragraph":
            p = doc.add_paragraph()
            _add_inline(p, node.get("children"))
            for run in p.runs:
                if run.font.name != "Cambria Math":
                    _set_run_font(run)
                if run.font.size is None:
                    run.font.size = Pt(14)
            # italic-only lines (editor notes): keep indent smaller optional
            if p.text.strip().startswith("*") and p.text.strip().endswith("*"):
                p.paragraph_format.first_line_indent = Cm(0)
                p.runs[0].italic = True if p.runs else None
            else:
                _paragraph_body_style(p)
            continue
        if t == "list":
            ordered = bool(node.get("attrs", {}).get("ordered"))
            for idx, item in enumerate(node.get("children") or [], start=1):
                if item.get("type") != "list_item":
                    continue
                prefix = f"{idx}. " if ordered else "• "
                p = doc.add_paragraph()
                p.paragraph_format.first_line_indent = Cm(0)
                p.paragraph_format.left_indent = Cm(1)
                p.paragraph_format.line_spacing_rule = WD_LINE_SPACING.MULTIPLE
                p.paragraph_format.line_spacing = 1.5
                pr = p.add_run(prefix)
                _set_run_font(pr)
                pr.bold = bool(ordered)
                _add_inline(p, _list_item_inline_children(item))
                for run in p.runs:
                    if run.font.name != "Cambria Math":
                        _set_run_font(run)
                    if run.font.size is None:
                        run.font.size = Pt(14)
            continue
        if t == "block_code":
            info = (node.get("attrs") or {}).get("info") or ""
            raw = (node.get("raw") or "").rstrip("\n")
            cap = doc.add_paragraph()
            cap.paragraph_format.first_line_indent = Cm(0)
            cr = cap.add_run(
                f"(Блок кода{f', язык: {info}' if info else ''}. "
                "Для Mermaid вставьте экспортированный рисунок из mermaid.live.)"
            )
            _set_run_font(cr)
            cr.italic = True
            cr.font.size = Pt(12)
            p = doc.add_paragraph()
            p.paragraph_format.first_line_indent = Cm(0)
            p.paragraph_format.left_indent = Cm(0.5)
            run = p.add_run(raw)
            _set_run_font(run, "Courier New")
            run.font.size = Pt(10)
            continue
        if t == "table":
            _add_table(doc, node)
            doc.add_paragraph()
            continue


def md_to_docx(src: Path, dst: Path) -> None:
    text = _latex_to_dollar_math(src.read_text(encoding="utf-8"))
    md = mistune.Markdown(plugins=[table, math])
    ast = md(text)
    if not isinstance(ast, list):
        ast = [ast]

    doc = Document()
    _configure_styles(doc)
    sect = doc.sections[0]
    sect.top_margin = Cm(2)
    sect.bottom_margin = Cm(2)
    sect.left_margin = Cm(3)
    sect.right_margin = Cm(1.5)

    _process_blocks(doc, ast)
    dst.parent.mkdir(parents=True, exist_ok=True)
    doc.save(str(dst))


def main() -> int:
    ap = argparse.ArgumentParser(description="Markdown → DOCX (VKR chapters)")
    ap.add_argument("src", type=Path, help="Path to .md file")
    ap.add_argument("dst", type=Path, nargs="?", help="Output .docx (default: same name as .md)")
    args = ap.parse_args()
    src: Path = args.src
    if not src.is_file():
        print(f"Not found: {src}", file=sys.stderr)
        return 1
    dst: Path = args.dst if args.dst else src.with_suffix(".docx")
    md_to_docx(src, dst)
    print(dst)
    return 0


if __name__ == "__main__":
    raise SystemExit(main())
